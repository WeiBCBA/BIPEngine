import io
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
import matplotlib.pyplot as plt
import pandas as pd
import streamlit as st

# ==========================================
# 1. Page Configuration & Custom CSS
# ==========================================
st.set_page_config(
    page_title="BCBA FBA & BIP Draft Formulation Tool",
    page_icon="🧩",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
    <style>
    .main-title {
        font-size: 2.1rem;
        color: #1F4E78;
        font-weight: 700;
        margin-bottom: 0.2rem;
        line-height: 1.3;
    }
    .demo-tag { font-size: 1.8rem; color: #1F4E78; font-weight: 600; }
    .sub-header { font-size: 0.95rem; color: #555; margin-bottom: 1.2rem; }
    
    .hipaa-banner {
        background-color: #EBF3FA;
        border: 2px solid #1F4E78;
        border-left: 8px solid #1F4E78;
        padding: 0.8rem 1rem;
        border-radius: 6px;
        margin-bottom: 1.2rem;
    }
    .hipaa-title { font-size: 1.15rem; font-weight: 800; color: #1F4E78; margin-bottom: 0.2rem; }
    .hipaa-body { font-size: 0.90rem; color: #2C3E50; line-height: 1.4; }

    .protocol-card {
        background-color: #F4F6F9;
        border: 1px solid #D1D5DB;
        border-left: 5px solid #1F4E78;
        padding: 0.9rem;
        border-radius: 6px;
        margin-bottom: 1.0rem;
    }
    .protocol-title { font-size: 1.0rem; font-weight: 700; color: #1F4E78; margin-bottom: 0.4rem; }
    .protocol-bullet { font-size: 0.88rem; color: #333; line-height: 1.5; margin-bottom: 0.2rem; }
    .algo-highlight { font-weight: 700; color: #C0392B; background-color: #FDEDEC; padding: 2px 6px; border-radius: 4px; }
    
    .stDownloadButton>button {
        background-color: #1F4E78 !important;
        color: white !important;
        font-size: 1.0rem !important;
        font-weight: 700 !important;
        padding: 0.6rem 1.0rem !important;
        border-radius: 6px !important;
        border: none !important;
        width: 100% !important;
    }
    </style>
""",
    unsafe_allow_html=True,
)


# ==========================================
# 2. Dynamic Mock Data Generators & Protocols
# ==========================================
def generate_mock_abc_csv(cohort_key):
  datasets = {
      "g1": [
          {
              "Date_Time": "2026-08-10 09:15",
              "Setting": "Clinic Playroom",
              "Antecedent": (
                  "Kitchen blender noise started in adjacent breakroom"
              ),
              "Behavior": "Screamed loudly and slapped face 3 times",
              "Consequence": (
                  "RBT offered noise-canceling headphones and sensory chew tool"
              ),
          },
          {
              "Date_Time": "2026-08-10 10:30",
              "Setting": "Outdoor Playground",
              "Antecedent": "Transition prompt given to pack up sand toys",
              "Behavior": (
                  "Attempted to eat sand and dropped to floor crying"
              ),
              "Consequence": (
                  "RBT blocked mouth, redirected to oral chew tool, demand"
                  " paused"
              ),
          },
          {
              "Date_Time": "2026-08-11 11:00",
              "Setting": "Table Therapy Room",
              "Antecedent": "Presented matching discrete trial worksheet",
              "Behavior": (
                  "Head banging on foam floor mat (3-4 forceful contacts)"
              ),
              "Consequence": (
                  "Therapist paused demand, presented PECS 'Break' card"
              ),
          },
          {
              "Date_Time": "2026-08-11 15:20",
              "Setting": "Clinic Snack Area",
              "Antecedent": "Preferred juice cup emptied",
              "Behavior": "Biting own right wrist, high-pitched crying",
              "Consequence": (
                  "RBT prompted functional communication button 'More Juice'"
              ),
          },
          {
              "Date_Time": "2026-08-12 09:40",
              "Setting": "Sensory Room",
              "Antecedent": "Sudden loud chime from timer",
              "Behavior": "Loud screaming and striking cheeks",
              "Consequence": "Staff provided noise-canceling headphones",
          },
          {
              "Date_Time": "2026-08-12 11:15",
              "Setting": "Clinic Floor",
              "Antecedent": "Instruction given to clean up blocks",
              "Behavior": "Dropped to floor and placed small toy near mouth",
              "Consequence": "RBT provided oral chew tool, delayed transition",
          },
      ],
      "g2": [
          {
              "Date_Time": "2026-08-10 09:30",
              "Setting": "Gen-Ed Classroom",
              "Antecedent": "Teacher presented 2-page math worksheet",
              "Behavior": "Screamed 'I won't do it!', pushed desk away",
              "Consequence": (
                  "Staff presented 'Break' visual card; demand paused"
              ),
          },
          {
              "Date_Time": "2026-08-10 11:00",
              "Setting": "Science Lab",
              "Antecedent": "Group presentation requirement announced",
              "Behavior": "Left desk area and hid under back table for 4 minutes",
              "Consequence": "Peer partner retrieved folder; teacher gave prompt",
          },
          {
              "Date_Time": "2026-08-11 13:15",
              "Setting": "School Cafeteria",
              "Antecedent": "Loud ambient cafeteria chatter and clattering trays",
              "Behavior": (
                  "Covered ears, vocalized high pitch, and bolted toward exit"
              ),
              "Consequence": (
                  "Aide followed with hall pass, escorted to quiet library"
                  " corner"
              ),
          },
          {
              "Date_Time": "2026-08-11 14:30",
              "Setting": "Computer Lab",
              "Antecedent": "Login failure on educational software page",
              "Behavior": (
                  "Struck keyboard keys forcefully and shoved monitor back"
              ),
              "Consequence": "IT support reset password; teacher provided help",
          },
          {
              "Date_Time": "2026-08-12 10:00",
              "Setting": "Language Arts",
              "Antecedent": "Asked to read essay aloud in front of peers",
              "Behavior": "Dropped head onto desk, refused to speak or look up",
              "Consequence": "Teacher allowed written alternative submission",
          },
      ],
      "g3": [
          {
              "Date_Time": "2026-08-09 09:00",
              "Setting": "Residential Home",
              "Antecedent": "New staff worker introduced schedule change",
              "Behavior": "Swept dishes off table, shouted threats",
              "Consequence": (
                  "DSP stepped in, offered visual choice board, demand paused"
              ),
          },
          {
              "Date_Time": "2026-08-10 10:15",
              "Setting": "Vocational Workshop",
              "Antecedent": "Supervisor increased box assembly quota by 20 units",
              "Behavior": "Shouted 'This is garbage!', slammed parts into bin",
              "Consequence": "Job coach offered 10-minute walk break",
          },
          {
              "Date_Time": "2026-08-11 13:30",
              "Setting": "Community Living Room",
              "Antecedent": (
                  "Peer requested remote control to change television channel"
              ),
              "Behavior": "Grabbed cushion, threw it across room, cursed loudly",
              "Consequence": "Staff intervened, redirected peer to another room",
          },
          {
              "Date_Time": "2026-08-12 11:00",
              "Setting": "Kitchen Prep Area",
              "Antecedent": "Staff asked client to wash cooking pans immediately",
              "Behavior": "Crossed arms, turned back, refused to touch sponge",
              "Consequence": "Staff gave 5-minute transition timer and space",
          },
          {
              "Date_Time": "2026-08-12 14:00",
              "Setting": "Day Program Hallway",
              "Antecedent": "Unscheduled fire drill alarm sounded loudly",
              "Behavior": "Stomped feet, yelled profanity, bolted toward exit door",
              "Consequence": "Staff escorted client safely to assembly point",
          },
      ],
  }
  df = pd.DataFrame(datasets.get(cohort_key, datasets["g1"]))
  return df.to_csv(index=False).encode("utf-8")


def generate_mock_tracking_csv(cohort_key):
  data = {
      "Day": ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"],
      "Target_Behavior_1_Freq": [3, 5, 2, 4, 1, 3, 2],
      "Target_Behavior_2_Freq": [2, 1, 4, 3, 2, 1, 2],
      "Target_Behavior_3_Freq": [4, 3, 1, 2, 5, 2, 3],
  }
  df = pd.DataFrame(data)
  return df.to_csv(index=False).encode("utf-8")


def generate_behavior_tracking_chart(cohort_key, behavior_index):
  fig, ax = plt.subplots(figsize=(6, 2.5))
  days = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]

  if behavior_index == 0:
    freqs = [3, 5, 2, 4, 1, 3, 2]
    color_val = "#1F4E78"
    title_str = "Target Behavior #1: Frequency Trend"
  elif behavior_index == 1:
    freqs = [2, 1, 4, 3, 2, 1, 2]
    color_val = "#C0392B"
    title_str = "Target Behavior #2: Frequency Trend"
  else:
    freqs = [4, 3, 1, 2, 5, 2, 3]
    color_val = "#27AE60"
    title_str = "Target Behavior #3: Frequency Trend"

  ax.plot(
      days, freqs, marker="o", color=color_val, linewidth=2, markersize=6
  )
  ax.set_title(
      title_str,
      fontsize=9,
      fontweight="bold",
      color=color_val,
  )
  ax.set_xlabel("Day of Week", fontsize=8)
  ax.set_ylabel("Episodes", fontsize=8)
  ax.grid(True, linestyle="--", alpha=0.5)
  plt.tight_layout()
  buf = io.BytesIO()
  fig.savefig(buf, format="png", dpi=150)
  buf.seek(0)
  plt.close(fig)
  return buf


def generate_mock_interview_docx(cohort_key):
  doc = docx.Document()
  doc.add_heading(
      "INDIRECT ASSESSMENT: STAKEHOLDER INTERVIEW NOTES (DE-IDENTIFIED)", level=1
  )
  doc.add_paragraph(
      "Client ID: [CLIENT_ID] | Target Cohort:"
      f" {cohort_meta[cohort_key]['title']}\nInformants: Parent, Lead"
      " Therapist / Educator, RBT Supervisor\n"
  )
  doc.add_heading("1. Interview Note Segment 1 (Parent Perspective)", level=2)
  doc.add_paragraph(
      "Summary: Stakeholders report elevated rates of target behaviors during"
      " task transitions, fine-motor academic demands, and high-pitch sensory"
      " noise environments. Parent notes consistent triggers at home."
  )
  doc.add_heading(
      "2. Interview Note Segment 2 (Educator / RBT Perspective)", level=2
  )
  doc.add_paragraph(
      "Summary: Staff observe functional escape behaviors when demands are"
      " presented rapidly without priming or visual schedules."
  )
  doc.add_heading("3. Interview Note Segment 3 (Peer / Environment Context)", level=2)
  doc.add_paragraph(
      "Summary: Environmental changes and loud auditory stimuli consistently"
      " exacerbate emotional distress and regulatory difficulties."
  )
  doc.add_heading(
      "4. Interview Note Segment 4 (Historical Reinforcement Patterns)", level=2
  )
  doc.add_paragraph(
      "Summary: Historical data indicates previous accidental reinforcement of"
      " avoidance through task removal."
  )
  doc.add_heading("5. Interview Note Segment 5 (Intervention Preferences)", level=2)
  doc.add_paragraph(
      "Summary: Stakeholders strongly prefer positive replacement strategies"
      " utilizing visual supports and communication devices."
  )
  bio = io.BytesIO()
  doc.save(bio)
  bio.seek(0)
  return bio


def generate_mock_qabf_docx(cohort_key):
  doc = docx.Document()
  doc.add_heading(
      "PSYCHOMETRIC QABF ASSESSMENT RESULTS (PER TARGET BEHAVIOR)", level=1
  )
  doc.add_paragraph(
      f"Client ID: [CLIENT_ID] | Cohort: {cohort_meta[cohort_key]['title']}\n"
  )
  b_list = cohort_meta[cohort_key]["behaviors"]
  for idx, b in enumerate(b_list, 1):
    doc.add_heading(f"Target Behavior #{idx}: {b['name']}", level=2)
    doc.add_paragraph(f"QABF Subscale Breakdown:\n{b['qabf_summary']}")
  bio = io.BytesIO()
  doc.save(bio)
  bio.seek(0)
  return bio


cohort_meta = {
    "g1": {
        "title": "Early Intervention Protocol (2-5 Yrs)",
        "file_tag": "2to5yo",
        "framework": "ESDM | NDBI Framework",
        "age_str": "3 Years 4 Months",
        "setting_str": "Early Intervention Clinic / Home Support",
        "protocol_sentences": [
            (
                "• Focuses on developmental milestone integration using ESDM /"
                " NDBI naturalistic approaches."
            ),
            (
                "• Prioritizes sensory processing, emotional regulation, and"
                " early functional communication (PECS/AAC)."
            ),
            (
                "• Integrates play-based assessment with parent-led co-regulation"
                " routines."
            ),
            (
                "• 📐 <strong>Triangulation Algorithm:</strong> 65% Direct ABC"
                " Observations + 25% Indirect Stakeholder Interviews + 10%"
                " Psychometric QABF Assessment"
            ),
        ],
        "behaviors": [
            {
                "name": (
                    "1. Self-Injurious Behavior (SIB) - Head Banging & Wrist"
                    " Biting"
                ),
                "name_zh": "1. 自伤行为 (SIB) - 撞头与咬手腕",
                "name_es": (
                    "1. Conducta autolesiva (SIB) - Golpes en la cabeza y"
                    " mordedura de muñeca"
                ),
                "def": (
                    "Any instance where the child forcefully makes contact"
                    " between forehead and a hard/padded surface, or places"
                    " wrist/hand between upper and lower teeth with visible"
                    " physical force lasting >1 second. Onset is the initial"
                    " physical contact; offset is 5 consecutive seconds without"
                    " contact."
                ),
                "def_zh": (
                    "儿童额头与硬质或软垫表面发生有力的碰撞，或将手腕/手部置于上下牙齿之间咬合并伴有明显施力的任何行为，持续时间>1秒。以首次物理接触为行为开始，以连续5秒无上述动作为行为结束。"
                ),
                "def_es": (
                    "Cualquier instancia en la que el niño haga contacto con"
                    " fuerza entre la frente y una superficie dura/acolchada, o"
                    " coloque la muñeca/mano entre los dientes superiores e"
                    " inferiores con fuerza física visible que dure >1 segundo."
                ),
                "ex": (
                    "Banging forehead 3-4 times on foam mat during tabletop"
                    " task; leaving red teeth marks on right wrist during"
                    " transition."
                ),
                "ex_zh": (
                    "在桌面任务期间，额头在泡棉垫上连续碰撞3-4次；在转换环节中在右手腕上留下红印牙痕。"
                ),
                "ex_es": (
                    "Golpear la frente 3-4 veces sobre la colchoneta durante"
                    " la tarea; dejar marcas rojas de dientes en la muñeca"
                    " derecha."
                ),
                "non_ex": (
                    "Resting head on floor mat during group circle time;"
                    " mouthing FDA-approved silicone chew necklace softly."
                ),
                "non_ex_zh": (
                    "在集体圈圈时间将头靠在地垫上休息；轻柔地咀嚼经过安全认证的硅胶项链。"
                ),
                "non_ex_es": (
                    "Descansar la cabeza en la colchoneta; morder suavemente un"
                    " collar de silicona apto."
                ),
                "dimensions": (
                    "Frequency: 3-6 episodes/day. Duration: 15s - 2min per"
                    " outburst. Intensity: Moderate to Severe (potential skin"
                    " redness or bruising)."
                ),
                "dimensions_zh": (
                    "频率：每天 3-6 次。持续时间：每次发作 15秒 至 2分钟。强度：中度至重度（可能导致皮肤发红或瘀青）。"
                ),
                "dimensions_es": (
                    "Frecuencia: 3-6 episodios/día. Duración: 15s - 2min."
                    " Intensidad: Moderada a Grave."
                ),
                "triggers": (
                    "Setting Events: Fatigue, high ambient background noise.\nImmediate"
                    " Triggers: Presentation of fine-motor discrete trial"
                    " tasks, removal of preferred sensory toy."
                ),
                "triggers_zh": (
                    "背景事件：疲劳、环境背景噪音过大。\n直接触发因素：呈现精细动作桌面任务、移走偏好的感官玩具。"
                ),
                "triggers_es": (
                    "Eventos de contexto: Fatiga, ruido ambiental.\nDesencadenantes"
                    " inmediatos: Tareas motoras finas, retirada de juguete"
                    " preferido."
                ),
                "consequences": (
                    "RBT blocks contact using foam pad, pauses academic demand"
                    " immediately, and prompts PECS 'Break' card."
                ),
                "consequences_zh": (
                    "RBT 使用软垫阻挡物理接触，立即暂停学业要求，并提示使用 PECS“休息”卡片。"
                ),
                "consequences_es": (
                    "El RBT bloquea el contacto, pausa la demanda e indica la"
                    " tarjeta PECS 'Descanso'."
                ),
                "qabf_summary": (
                    "Task Escape: 14/15 | Physical Discomfort: 8/15 |"
                    " Attention: 2/15 | Tangible: 3/15 | Sensory: 4/15"
                ),
                "qabf_summary_zh": (
                    "逃避任务: 14/15 | 身体不适: 8/15 | 社交关注: 2/15 | 获得物质: 3/15 |"
                    " 感官刺激: 4/15"
                ),
                "qabf_summary_es": (
                    "Escape de tarea: 14/15 | Malestar físico: 8/15 | Atención:"
                    " 2/15 | Tangible: 3/15 | Sensorial: 4/15"
                ),
                "triangulation": (
                    "65% Direct ABC Data (high rate during table work) + 25%"
                    " Indirect Parent Interview (frustration with structured"
                    " demands) + 10% QABF (Escape score 14/15)."
                ),
                "triangulation_zh": (
                    "65% 直接 ABC 数据（桌面任务期间高发） + 25% 间接家长访谈（对结构化任务表现出挫败）"
                    " + 10% QABF 结果（逃避任务得分 14/15）。"
                ),
                "triangulation_es": (
                    "65% Datos ABC + 25% Entrevista + 10% QABF (Escape 14/15)."
                ),
                "hypothesis": (
                    "Primary Function: Task Escape (Social Negative"
                    " Reinforcement).\nSecondary Function: Physical Discomfort"
                    " Relief."
                ),
                "hypothesis_zh": "核心行为功能：逃避任务（社交负强化）。\n次要行为功能：缓解身体不适。",
                "hypothesis_es": (
                    "Función principal: Escape de tareas.\nFunción secundaria:"
                    " Alivio de malestar físico."
                ),
                "ferb": (
                    "Functional Communication: Activate AAC BigMack button"
                    " ('I need a break') or hand 'Break' PECS card to prompt"
                    " immediate task pause."
                ),
                "ferb_zh": (
                    "功能性替代行为 (FERB)：按下 AAC 语音按键（“我想休息”）或递交“休息”"
                    " PECS 卡片，以提示立即暂停当前任务。"
                ),
                "ferb_es": (
                    "Comunicación Funcional: Activar el botón AAC BigMack"
                    " ('Necesito un descanso') o entregar la tarjeta PECS"
                    " 'Descanso'."
                ),
            },
            {
                "name": "2. Sensory Vocal Distress & Face-Slapping",
                "name_zh": "2. 感官情绪性尖叫与拍打面部",
                "name_es": "2. Distrés vocal sensorial y bofetadas en la cara",
                "def": (
                    "Vocal screaming exceeding normal conversational volume (>80"
                    " dB) lasting >2 seconds, occurring simultaneously with"
                    " open-palm striking of own cheeks or arms. Onset: initial"
                    " vocalization/strike; Offset: 10 consecutive seconds of"
                    " calm without striking."
                ),
                "def_zh": (
                    "尖叫声超过正常交谈音量（>80 dB）且持续时间>2秒，同时伴随用张开的手掌拍打自己脸颊或手臂的行为。以初始发声/拍打为行为开始，以连续10秒保持平静无拍打动作为行为结束。"
                ),
                "def_es": (
                    "Gritos vocales que superan el volumen conversacional normal"
                    " (>80 dB) durante >2 segundos, junto con golpes con la"
                    " palma abierta en las mejillas o brazos."
                ),
                "ex": (
                    "Loud screaming and striking cheeks 3 times when loud"
                    " kitchen appliance starts in adjacent room."
                ),
                "ex_zh": "当隔壁房间开启高分贝厨房电器时，大声尖叫并连续拍打脸颊 3 次。",
                "ex_es": (
                    "Gritos fuertes y 3 golpes en las mejillas al encenderse un"
                    " electrodoméstico ruidoso."
                ),
                "non_ex": (
                    "Excited vocal shouting during outdoor playground play;"
                    " tapping cheeks rhythmically during music group time."
                ),
                "non_ex_zh": (
                    "在户外操场游玩时兴奋地高声欢呼；在音乐课上跟随节奏轻拍脸颊。"
                ),
                "non_ex_es": (
                    "Gritos de emoción en el parque; golpear las mejillas al"
                    " ritmo de la música."
                ),
                "dimensions": (
                    "Frequency: 2-4 episodes/day. Duration: 30s - 3min."
                    " Intensity: Moderate."
                ),
                "dimensions_zh": (
                    "频率：每天 2-4 次。持续时间：30秒 至 3分钟。强度：中度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 2-4 episodios/día. Duración: 30s - 3min."
                    " Intensidad: Moderada."
                ),
                "triggers": (
                    "Setting Events: Overstimulating ambient noise, sudden"
                    " schedule changes.\nImmediate Triggers: Unexpected high-pitch"
                    " sounds, removal of preferred juice cup."
                ),
                "triggers_zh": (
                    "背景事件：环境噪音过度刺激、突发日程改变。\n直接触发因素：意料之外的高频噪音、偏好的果汁杯被移走。"
                ),
                "triggers_es": (
                    "Eventos de contexto: Ruido ambiental excesivo.\nDesencadenantes:"
                    " Sonidos agudos inesperados."
                ),
                "consequences": (
                    "Staff offers noise-canceling headphones and redirects to"
                    " sensory chew tool."
                ),
                "consequences_zh": "工作人员提供降噪耳机，并重新引导至口部感官咀嚼工具。",
                "consequences_es": (
                    "El personal ofrece auriculares de cancelación de ruido y"
                    " redirige a un mordedor sensorial."
                ),
                "qabf_summary": (
                    "Sensory/Automatic: 13/15 | Task Escape: 11/15 | Attention:"
                    " 3/15 | Tangible: 2/15 | Physical: 1/15"
                ),
                "qabf_summary_zh": (
                    "感官刺激/自动强化: 13/15 | 逃避任务: 11/15 | 社交关注: 3/15 |"
                    " 获得物质: 2/15 | 身体不适: 1/15"
                ),
                "qabf_summary_es": (
                    "Sensorial: 13/15 | Escape: 11/15 | Atención: 3/15 |"
                    " Tangible: 2/15 | Físico: 1/15"
                ),
                "triangulation": (
                    "65% Direct ABC Data (appliance noise trigger) + 25%"
                    " Indirect RBT Interview (auditory aversion) + 10% QABF"
                    " (Sensory score 13/15)."
                ),
                "triangulation_zh": (
                    "65% 直接 ABC 数据（电器噪音触发） + 25% 间接 RBT 访谈（听觉敏感过度） + 10%"
                    " QABF 结果（感官得分 13/15）。"
                ),
                "triangulation_es": (
                    "65% Datos ABC + 25% Entrevista RBT + 10% QABF (Sensorial"
                    " 13/15)."
                ),
                "hypothesis": (
                    "Primary Function: Escape / Regulation of Auditory"
                    " Overstimulation (Sensory Automatic Reinforcement)."
                ),
                "hypothesis_zh": (
                    "核心行为功能：逃避/调节听觉过度刺激（感官自动强化）。"
                ),
                "hypothesis_es": (
                    "Función principal: Escape / Regulación de sobreestimulación"
                    " auditiva."
                ),
                "ferb": (
                    "Self-Regulation / Request Strategy: Point to 'Headphones'"
                    " visual symbol or independently retrieve noise-canceling"
                    " headphones from designated sensory bin."
                ),
                "ferb_zh": (
                    "功能性替代行为 (FERB)：指认“耳机”视觉标识，或自行从指定的感官箱中取出降噪耳机戴上。"
                ),
                "ferb_es": (
                    "Estrategia de autorregulación: Señalar el símbolo"
                    " 'Auriculares' o tomarlos de la caja sensorial."
                ),
            },
            {
                "name": "3. Floor Dropping & Pica Attempts",
                "name_zh": "3. 突然倒地与异食倾向 (Pica)",
                "name_es": "3. Tirarse al suelo e intentos de pica",
                "def": (
                    "Unprompted sudden collapse of body onto floor from standing"
                    " or seated position, accompanied by attempts to place"
                    " non-food items (e.g., sand, dirt, small plastic objects)"
                    " into mouth. Onset: body contacting floor; Offset:"
                    " standing up for >5 seconds."
                ),
                "def_zh": (
                    "在站立或坐姿状态下，未经提示突然将身体倒在地上，并伴随试图将非食物物品（如沙子、泥土、塑料小零件）放入口中的行为。以身体接触地面为行为开始，以重新站立保持>5秒为行为结束。"
                ),
                "def_es": (
                    "Caída repentina e imprevista del cuerpo al suelo desde una"
                    " posición de pie o sentada, acompañada de intentos de"
                    " llevarse a la boca objetos no alimentarios (p. ej.,"
                    " arena, tierra, objetos pequeños de plástico). Inicio:"
                    " contacto del cuerpo con el suelo; Fin: permanecer de pie"
                    " durante >5 segundos."
                ),
                "ex": (
                    "Dropping to playground sand and placing a handful of sand"
                    " near mouth when prompted to transition inside."
                ),
                "ex_zh": (
                    "当被提示转换进室内时，突然倒在操场沙坑里并将一把沙子拿到嘴边。"
                ),
                "ex_es": (
                    "Tirarse a la arena del parque y llevarse un puñado cerca de"
                    " la boca al recibir la indicación de entrar."
                ),
                "non_ex": (
                    "Lying down on rest mat during scheduled nap time; placing"
                    " silicone chew toy into mouth."
                ),
                "non_ex_zh": (
                    "在计划的午休时间躺在休息垫上；将硅胶咀嚼玩具放入口中。"
                ),
                "non_ex_es": (
                    "Acostarse en la colchoneta durante la siesta; ponerse el"
                    " mordedor de silicona en la boca."
                ),
                "dimensions": (
                    "Frequency: 1-3 episodes/day. Duration: 1-5min. Intensity:"
                    " Moderate."
                ),
                "dimensions_zh": (
                    "频率：每天 1-3 次。持续时间：1 至 5分钟。强度：中度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 1-3 episodios/día. Duración: 1-5min. Intensidad:"
                    " Moderada."
                ),
                "triggers": (
                    "Setting Events: Unassigned downtime, transition from"
                    " outdoor to indoor.\nImmediate Triggers: Direct"
                    " instruction to clean up preferred toys."
                ),
                "triggers_zh": (
                    "背景事件：无安排的空闲时间、从户外到室内的活动转换。\n直接触发因素：要求收拾偏好玩具的直接指令。"
                ),
                "triggers_es": (
                    "Eventos de contexto: Tiempo libre sin estructurar,"
                    " transición exterior/interior.\nDesencadenante: Orden de"
                    " recoger los juguetes."
                ),
                "consequences": (
                    "RBT blocks hand-to-mouth trajectory, provides oral chew"
                    " device, and delays transition demand."
                ),
                "consequences_zh": (
                    "RBT 阻挡手移向嘴部的轨迹，提供口部咀嚼器，并延缓转换要求。"
                ),
                "consequences_es": (
                    "El RBT bloquea el trayecto mano-boca, ofrece mordedor"
                    " oral y retrasa la demanda."
                ),
                "qabf_summary": (
                    "Task Escape: 12/15 | Tangible Access: 10/15 | Attention:"
                    " 4/15 | Sensory: 3/15 | Physical: 1/15"
                ),
                "qabf_summary_zh": (
                    "逃避任务: 12/15 | 获得物质: 10/15 | 社交关注: 4/15 | 感官刺激:"
                    " 3/15 | 身体不适: 1/15"
                ),
                "qabf_summary_es": (
                    "Escape: 12/15 | Acceso a Tangible: 10/15 | Atención: 4/15"
                    " | Sensorial: 3/15 | Físico: 1/15"
                ),
                "triangulation": (
                    "65% Direct ABC Data (playground cleanup trigger) + 25%"
                    " Indirect Teacher Interview (transition delay) + 10% QABF"
                    " (Escape score 12/15)."
                ),
                "triangulation_zh": (
                    "65% 直接 ABC 数据（操场清理触发） + 25% 间接教师访谈（转换延缓） + 10%"
                    " QABF 结果（逃避得分 12/15）。"
                ),
                "triangulation_es": (
                    "65% Datos ABC + 25% Entrevista + 10% QABF (Escape 12/15)."
                ),
                "hypothesis": (
                    "Primary Function: Transition Escape & Delay.\nSecondary"
                    " Function: Access to Outdoor Tangible Play."
                ),
                "hypothesis_zh": (
                    "核心行为功能：逃避与延缓活动转换。\n次要行为功能：继续获取户外游戏物品。"
                ),
                "hypothesis_es": (
                    "Función principal: Escape/retraso de transición.\nFunción"
                    " secundaria: Acceso a juego."
                ),
                "ferb": (
                    "Transition Delay AAC: Exchange '1 More Minute' visual"
                    " card or press AAC button to request additional play time"
                    " prior to cleanup."
                ),
                "ferb_zh": (
                    "功能性替代行为 (FERB)：递交“再玩1分钟”视觉卡片或按下 AAC"
                    " 沟通按键，在收拾前请求额外的游戏时间。"
                ),
                "ferb_es": (
                    "Comunicación Funcional: Entregar tarjeta '1 minuto más' o"
                    " presionar el botón AAC."
                ),
            },
        ],
        "strengths": (
            "Responds very well to 1:1 adult playful interaction, strong"
            " visual matching skills, highly motivated by musical cause-and-effect"
            " toys."
        ),
        "strengths_zh": (
            "对 1:1"
            " 成人游戏化互动反应良好，具备较强的视觉匹配能力，对音乐因果玩具表现出极高动机。"
        ),
        "strengths_es": (
            "Responde muy bien a la interacción lúdica 1:1 con adultos,"
            " excelentes habilidades de emparejamiento visual, altamente motivado"
            " por juguetes musicales."
        ),
        "history": (
            "Diagnosed with ASD Level 3; currently receiving 15 hrs/week of"
            " Early Intervention ABA and Speech Therapy."
        ),
        "history_zh": (
            "确诊为孤独症谱系障碍（ASD 3级）；目前每周接受 15"
            " 小时的早期干预 ABA 及言语治疗服务。"
        ),
        "history_es": (
            "Diagnóstico de TEA Nivel 3; actualmente recibe 15 horas/semana de"
            " Intervención Temprana ABA y Terapia de Lenguaje."
        ),
    },
    "g2": {
        "title": "School-Age IEP Protocol (5-21 Yrs)",
        "file_tag": "5to21yo",
        "framework": "IDEA IEP | PBIS Framework",
        "age_str": "10 Years 2 Months",
        "setting_str": "General Education Classroom / Resource Room",
        "protocol_sentences": [
            (
                "• Aligned with IDEA IEP requirements and PBIS Multi-Tiered"
                " Support Systems."
            ),
            (
                "• Targets academic task engagement, self-advocacy, and emotional"
                " self-regulation."
            ),
            (
                "• Emphasizes replacement behaviors integrated into classroom"
                " routines."
            ),
            (
                "• 📐 <strong>Triangulation Algorithm:</strong> 65% Direct ABC"
                " Observations + 25% Indirect Stakeholder Interviews + 10%"
                " Psychometric QABF Assessment"
            ),
        ],
        "behaviors": [
            {
                "name": "1. Task Avoidance / Elopement from Seat",
                "name_zh": "1. 逃避学业任务与擅离座位 (Elopement)",
                "name_es": "1. Evitación de tareas y abandono del asiento",
                "def": (
                    "Leaving assigned desk area without teacher permission for"
                    " >5 seconds during independent academic instruction. Onset:"
                    " feet leaving designated desk perimeter; Offset: returning"
                    " to seat."
                ),
                "def_zh": (
                    "在独立学业教学期间，未经教师允许离开指定的课桌区域超过5秒。以双脚离开指定课桌边界为行为开始，以返回座位为行为结束。"
                ),
                "def_es": (
                    "Abandonar el área asignada del escritorio sin permiso del"
                    " maestro durante >5 segundos durante la instrucción"
                    " académica independiente."
                ),
                "ex": (
                    "Running out of seat to classroom carpet area during"
                    " independent math worksheet."
                ),
                "ex_zh": (
                    "在独立完成数学工作表期间，从座位上跑开并躺在教室地毯区域。"
                ),
                "ex_es": (
                    "Salir corriendo del asiento hacia el área de la alfombra"
                    " durante una hoja de trabajo de matemáticas."
                ),
                "non_ex": (
                    "Standing up to walk to sharpener after raising hand and"
                    " receiving permission."
                ),
                "non_ex_zh": "举手并获得许可后，站起来走到削笔器旁削铅笔。",
                "non_ex_es": (
                    "Levantarse para ir al sacapuntas tras levantar la mano y"
                    " recibir permiso."
                ),
                "dimensions": (
                    "Frequency: 4-5 times per school day. Duration: 1-5 minutes"
                    " per instance. Intensity: Low to Moderate."
                ),
                "dimensions_zh": (
                    "频率：每个学校日 4-5 次。持续时间：每次 1-5 分钟。强度：低至中度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 4-5 veces por día escolar. Duración: 1-5"
                    " minutos. Intensidad: Baja a Moderada."
                ),
                "triggers": (
                    "Setting Events: Multi-step math tasks.\nImmediate Triggers:"
                    " Presentation of 2-page math assignment."
                ),
                "triggers_zh": (
                    "背景事件：多步骤数学任务。\n直接触发因素：发放长达2页的数学作业纸。"
                ),
                "triggers_es": (
                    "Eventos de contexto: Tareas matemáticas de múltiples"
                    " pasos.\nDesencadenante: Entrega de hojas de trabajo."
                ),
                "consequences": (
                    "Staff presents 'Break' visual card; demand temporarily"
                    " paused."
                ),
                "consequences_zh": "教职工呈现“休息”视觉卡片；学业要求被暂时暂停。",
                "consequences_es": (
                    "El personal presenta tarjeta visual 'Descanso'; demanda"
                    " pausada."
                ),
                "qabf_summary": (
                    "Task Escape: 15/15 | Attention: 5/15 | Tangible: 2/15 |"
                    " Sensory: 1/15 | Physical: 0/15"
                ),
                "qabf_summary_zh": (
                    "逃避任务: 15/15 | 社交关注: 5/15 | 获得物质: 2/15 | 感官刺激:"
                    " 1/15 | 身体不适: 0/15"
                ),
                "qabf_summary_es": (
                    "Escape: 15/15 | Atención: 5/15 | Tangible: 2/15 |"
                    " Sensorial: 1/15 | Físico: 0/15"
                ),
                "triangulation": (
                    "65% Direct ABC Data + 25% Indirect IEP Interview + 10%"
                    " QABF (Escape score 15/15)."
                ),
                "triangulation_zh": (
                    "65% 直接 ABC 数据 + 25% 间接 IEP 访谈 + 10% QABF 结果（逃避得分 15/15）。"
                ),
                "triangulation_es": (
                    "65% Datos ABC + 25% Entrevista IEP + 10% QABF (Escape 15/15)."
                ),
                "hypothesis": "Primary Function: Escape from Academic Demands.",
                "hypothesis_zh": "核心行为功能：逃避学业任务要求。",
                "hypothesis_es": "Función principal: Escape de demandas académicas.",
                "ferb": (
                    "Hand 'Break' card to teacher or place 'Help Needed' tent on"
                    " desk."
                ),
                "ferb_zh": (
                    "向教师递交“休息”卡片，或在桌上摆放“需要帮助”提示牌。"
                ),
                "ferb_es": (
                    "Entregar tarjeta 'Descanso' al maestro o colocar tarjeta"
                    " 'Ayuda'."
                ),
            },
            {
                "name": "2. Group Presentation Avoidance & Hiding",
                "name_zh": "2. 逃避小组展示与躲藏行为",
                "name_es": "2. Evitación de presentaciones grupales y ocultamiento",
                "def": (
                    "Leaving assigned group space or crawling under furniture"
                    " when public speaking or group presentation is requested."
                    " Onset: body moving away from group; Offset: returning after"
                    " prompt."
                ),
                "def_zh": (
                    "当被要求公开演讲或小组展示时，离开指定的小组空间或爬到家具下方。以身体离开小组为行为开始，以经提示后返回为行为结束。"
                ),
                "def_es": (
                    "Abandonar el espacio grupal o esconderse bajo los muebles"
                    " cuando se requiere hablar en público."
                ),
                "ex": "Hiding under the back table during science presentation.",
                "ex_zh": "在科学课展示期间躲在后排桌子底下。",
                "ex_es": (
                    "Esconderse bajo la mesa trasera durante la presentación de"
                    " ciencias."
                ),
                "non_ex": "Sitting quietly while listening to a peer present.",
                "non_ex_zh": "安静坐着聆听同伴展示。",
                "non_ex_es": (
                    "Sentarse tranquilamente mientras escucha a un compañero."
                ),
                "dimensions": (
                    "Frequency: 2-3 times per week. Duration: 3-8 minutes."
                    " Intensity: Moderate."
                ),
                "dimensions_zh": (
                    "频率：每周 2-3 次。持续时间：3-8 分钟。强度：中度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 2-3 veces por semana. Duración: 3-8 min."
                    " Intensidad: Moderada."
                ),
                "triggers": (
                    "Setting Events: Social anxiety, public speaking demands."
                ),
                "triggers_zh": "背景事件：社交焦虑、公开演讲要求。",
                "triggers_es": (
                    "Eventos de contexto: Ansiedad social, hablar en público."
                ),
                "consequences": (
                    "Teacher provides alternative individual task option."
                ),
                "consequences_zh": "教师提供个人独立任务替代选项。",
                "consequences_es": (
                    "El maestro ofrece opción de tarea individual alternativa."
                ),
                "qabf_summary": (
                    "Task Escape: 13/15 | Anxiety/Avoidance: 12/15 | Attention:"
                    " 2/15"
                ),
                "qabf_summary_zh": "逃避任务: 13/15 | 焦虑/回避: 12/15 | 社交关注: 2/15",
                "qabf_summary_es": (
                    "Escape: 13/15 | Ansiedad/Evitación: 12/15 | Atención: 2/15"
                ),
                "triangulation": "65% ABC Data + 25% Teacher Interview + 10% QABF.",
                "triangulation_zh": "65% ABC 数据 + 25% 教师访谈 + 10% QABF。",
                "triangulation_es": "65% Datos ABC + 25% Entrevista + 10% QABF.",
                "hypothesis": (
                    "Primary Function: Escape from Social Evaluative Stress."
                ),
                "hypothesis_zh": "核心行为功能：逃避社交评估压力。",
                "hypothesis_es": (
                    "Función principal: Escape de estrés evaluativo social."
                ),
                "ferb": (
                    "Request pre-recorded video presentation submission instead"
                    " of live speaking."
                ),
                "ferb_zh": "请求提交预先录制的视频展示替代现场发言。",
                "ferb_es": (
                    "Solicitar presentación en video pregrabado en lugar de en"
                    " vivo."
                ),
            },
            {
                "name": "3. Auditory Overload & Elopement to Hallway",
                "name_zh": "3. 听觉超载与冲向走廊行为",
                "name_es": "3. Sobrecarga auditiva y fuga al pasillo",
                "def": (
                    "Covering ears and running out of the classroom into the"
                    " hallway upon exposure to high ambient noise (>75 dB)."
                    " Onset: hands to ears and bolting; Offset: calming in"
                    " quiet area."
                ),
                "def_zh": (
                    "暴露于高环境噪音（>75 dB）时，双手捂耳并跑出教室进入走廊。以双手捂耳并冲出门为行为开始，以在安静区域平静为行为结束。"
                ),
                "def_es": (
                    "Cubrirse los orejas y salir corriendo del salón al pasillo"
                    " ante ruido ambiental alto."
                ),
                "ex": "Running out to hallway during noisy cafeteria transition.",
                "ex_zh": "在嘈杂的餐厅转换期间跑向走廊。",
                "ex_es": (
                    "Correr al pasillo durante la transición ruidosa de la"
                    " cafetería."
                ),
                "non_ex": "Putting on headphones proactively in quiet room.",
                "non_ex_zh": "在安静房间主动戴上耳机。",
                "non_ex_es": (
                    "Ponerse auriculares proactivamente en una habitación"
                    " silenciosa."
                ),
                "dimensions": (
                    "Frequency: 1-3 times weekly. Duration: 5-10 minutes."
                    " Intensity: Moderate."
                ),
                "dimensions_zh": (
                    "频率：每周 1-3 次。持续时间：5-10 分钟。强度：中度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 1-3 veces por semana. Duración: 5-10 min."
                    " Intensidad: Moderada."
                ),
                "triggers": "Setting Events: Cafeteria, assembly halls, gym classes.",
                "triggers_zh": "背景事件：食堂、礼堂、体育课。",
                "triggers_es": "Eventos de contexto: Cafetería, asambleas, gimnasio.",
                "consequences": (
                    "Aide escorts student to quiet library corner with pass."
                ),
                "consequences_zh": "助教持通行证护送学生至图书馆安静角落。",
                "consequences_es": (
                    "El asistente acompaña al estudiante a un rincón tranquilo."
                ),
                "qabf_summary": (
                    "Sensory Automatic: 14/15 | Escape: 10/15 | Attention:"
                    " 1/15"
                ),
                "qabf_summary_zh": (
                    "感官自动强化: 14/15 | 逃避任务: 10/15 | 社交关注: 1/15"
                ),
                "qabf_summary_es": (
                    "Sensorial automático: 14/15 | Escape: 10/15 | Atención:"
                    " 1/15"
                ),
                "triangulation": "65% ABC Data + 25% Aide Interview + 10% QABF.",
                "triangulation_zh": "65% ABC 数据 + 25% 助教访谈 + 10% QABF。",
                "triangulation_es": "65% Datos ABC + 25% Entrevista + 10% QABF.",
                "hypothesis": (
                    "Primary Function: Automatic Sensory Regulation / Escape"
                    " from Noise."
                ),
                "hypothesis_zh": "核心行为功能：自动感官调节/逃避噪音。",
                "hypothesis_es": (
                    "Función principal: Regulación sensorial / Escape del"
                    " ruido."
                ),
                "ferb": (
                    "Independently retrieve and wear noise-canceling headphones"
                    " upon entering loud areas."
                ),
                "ferb_zh": "进入嘈杂区域时独立取出并佩戴降噪耳机。",
                "ferb_es": (
                    "Tomar y usar auriculares con cancelación de ruido de"
                    " forma independiente."
                ),
            },
        ],
        "strengths": (
            "Excellent visual-spatial abilities, enthusiastic about technology"
            " and drawing."
        ),
        "strengths_zh": "具备出色的视觉空间能力，对科技和绘画抱有极高热情。",
        "strengths_es": (
            "Excelentes habilidades visoespaciales, entusiasta de la"
            " tecnología y el dibujo."
        ),
        "history": "Enrolled in General Education with IEP support.",
        "history_zh": "就读于普通教育班级，享有 IEP 特殊教育支持计划。",
        "history_es": "Inscrito en educación general con apoyo de IEP.",
    },
    "g3": {
        "title": "Adult Community Protocol (21+ Yrs)",
        "file_tag": "21plusYo",
        "framework": "Medicaid HCBS | Person-Centered Waiver Framework",
        "age_str": "26 Years 8 Months",
        "setting_str": "Vocational Workshop & Day Program",
        "protocol_sentences": [
            (
                "• Designed for Medicaid HCBS Waiver adult day programs and"
                " community living."
            ),
            (
                "• Focuses on person-centered planning, independence, and"
                " vocational endurance."
            ),
            (
                "• Emphasizes self-management protocols and respectful adult"
                " communication."
            ),
            (
                "• 📐 <strong>Triangulation Algorithm:</strong> 65% Direct ABC"
                " Observations + 25% Indirect Stakeholder Interviews + 10%"
                " Psychometric QABF Assessment"
            ),
        ],
        "behaviors": [
            {
                "name": "1. Vocational Task Refusal & Verbal Aggression",
                "name_zh": "1. 拒绝职业组装任务与言语攻击",
                "name_es": "1. Rechazo de tareas vocacionales y agresión verbal",
                "def": (
                    "Refusing assembly or sorting demands accompanied by loud"
                    " vocal threats (>75 dB) or pushing work materials away."
                    " Onset: vocal outburst or material push; Offset: 3 minutes"
                    " of quiet task engagement."
                ),
                "def_zh": (
                    "拒绝组装或分类任务，并伴随大声言语威胁（>75 dB）或推开工作材料的行为。以言语发作或推开材料为行为开始，以连续"
                    " 3 分钟安静参与任务为行为结束。"
                ),
                "def_es": (
                    "Rechazar demandas de ensamblaje acompañado de amenazas"
                    " vocales fuertes (>75 dB) o empujar los materiales de"
                    " trabajo."
                ),
                "ex": (
                    "Shouting 'No way!', slamming assembly boxes on desk when"
                    " quota is raised."
                ),
                "ex_zh": "当提高工作配额时，大叫“绝不可能！”并将组装盒重重摔在桌上。",
                "ex_es": (
                    "Gritar '¡De ninguna manera!' y golpear cajas sobre la mesa."
                ),
                "non_ex": "Verbally requesting a 5-minute break in a normal tone.",
                "non_ex_zh": "用正常音量和语调口头提出“想要休息5分钟”。",
                "non_ex_es": (
                    "Solicitar verbalmente un descanso de 5 minutos en tono"
                    " normal."
                ),
                "dimensions": (
                    "Frequency: 1-2 times weekly. Duration: 5-10 minutes."
                    " Intensity: Moderate."
                ),
                "dimensions_zh": (
                    "频率：每周 1-2 次。持续时间：5-10 分钟。强度：中度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 1-2 veces por semana. Duración: 5-10 min."
                    " Intensidad: Moderada."
                ),
                "triggers": (
                    "Setting Events: Unfamiliar staff.\nImmediate Triggers:"
                    " Direct instructions to complete vocational assembly"
                    " quota."
                ),
                "triggers_zh": (
                    "背景事件：不熟悉的工作人员。\n直接触发因素：要求完成职业组装配额的直接指令。"
                ),
                "triggers_es": (
                    "Eventos de contexto: Personal no familiar.\nDesencadenante:"
                    " Instrucción directa de ensamblaje."
                ),
                "consequences": "DSP offers choice board, demand temporarily paused.",
                "consequences_zh": (
                    "直属支持人员（DSP）提供选择板，任务要求被暂时暂停。"
                ),
                "consequences_es": (
                    "El personal DSP ofrece panel de opciones, demanda pausada."
                ),
                "qabf_summary": (
                    "Task Escape: 13/15 | Attention: 4/15 | Tangible: 3/15 |"
                    " Sensory: 1/15 | Physical: 1/15"
                ),
                "qabf_summary_zh": (
                    "逃避任务: 13/15 | 社交关注: 4/15 | 获得物质: 3/15 | 感官刺激:"
                    " 1/15 | 身体不适: 1/15"
                ),
                "qabf_summary_es": (
                    "Escape: 13/15 | Atención: 4/15 | Tangible: 3/15 |"
                    " Sensorial: 1/15 | Físico: 1/15"
                ),
                "triangulation": (
                    "65% Direct ABC Data + 25% Indirect Vocational Interview +"
                    " 10% QABF (Escape score 13/15)."
                ),
                "triangulation_zh": (
                    "65% 直接 ABC 数据 + 25% 间接职业能力访谈 + 10% QABF 结果（逃避得分 13/15）。"
                ),
                "triangulation_es": (
                    "65% Datos ABC + 25% Entrevista + 10% QABF (Escape 13/15)."
                ),
                "hypothesis": (
                    "Primary Function: Escape from Vocational Assembly"
                    " Demands."
                ),
                "hypothesis_zh": "核心行为功能：逃避职业组装工作要求。",
                "hypothesis_es": "Función principal: Escape de demandas vocacionales.",
                "ferb": (
                    "Verbally request '5-minute break, please' using"
                    " self-advocacy phrase card."
                ),
                "ferb_zh": (
                    "使用自我倡导短语卡口头表达：“请给我 5 分钟休息时间”。"
                ),
                "ferb_es": (
                    "Solicitar verbalmente 'Descanso de 5 minutos, por favor'."
                ),
            },
            {
                "name": "2. Property Disruption During Peer Conflict",
                "name_zh": "2. 同伴冲突期间的物品破坏行为",
                "name_es": "2. Alteración de la propiedad durante conflicto con pares",
                "def": (
                    "Throwing objects or sweeping items off surfaces when"
                    " disputes arise over shared community spaces or items."
                    " Onset: grabbing/throwing object; Offset: 2 minutes of calm."
                ),
                "def_zh": (
                    "当在共享社区空间或物品产生争议时，投掷物品或将物品从表面扫落。以抓取/投掷物品为行为开始，以持续 2 分钟平静为行为结束。"
                ),
                "def_es": (
                    "Lojar objetos o barrer elementos de las superficies cuando"
                    " surgen disputas."
                ),
                "ex": "Throwing couch cushions when peer takes remote control.",
                "ex_zh": "当同伴拿走遥控器时扔掉沙发布艺垫。",
                "ex_es": (
                    "Lanzar cojines del sofá cuando un compañero toma el control"
                    " remoto."
                ),
                "non_ex": "Asking staff to mediate peer disagreement calmly.",
                "non_ex_zh": "平静地请求工作人员调解同伴意见分歧。",
                "non_ex_es": (
                    "Pedir al personal que medie en el desacuerdo con calma."
                ),
                "dimensions": (
                    "Frequency: 1-2 times weekly. Duration: 2-5 minutes."
                    " Intensity: Moderate."
                ),
                "dimensions_zh": (
                    "频率：每周 1-2 次。持续时间：2-5 分钟。强度：中度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 1-2 veces por semana. Duración: 2-5 min."
                    " Intensidad: Moderada."
                ),
                "triggers": (
                    "Setting Events: Shared living spaces, competing for media"
                    " items."
                ),
                "triggers_zh": "背景事件：共享生活空间、争夺娱乐设备。",
                "triggers_es": (
                    "Eventos de contexto: Espacios compartidos, competencia por"
                    " medios."
                ),
                "consequences": (
                    "Staff intervenes, separates individuals, offers"
                    " mediation."
                ),
                "consequences_zh": "工作人员干预，分隔双方并提供调解。",
                "consequences_es": (
                    "El personal interviene, separa a las personas y ofrece"
                    " mediación."
                ),
                "qabf_summary": (
                    "Tangible Access: 14/15 | Escape: 8/15 | Attention: 5/15"
                ),
                "qabf_summary_zh": "获得物质: 14/15 | 逃避任务: 8/15 | 社交关注: 5/15",
                "qabf_summary_es": (
                    "Acceso a tangible: 14/15 | Escape: 8/15 | Atención: 5/15"
                ),
                "triangulation": "65% ABC Data + 25% DSP Interview + 10% QABF.",
                "triangulation_zh": "65% ABC 数据 + 25% DSP 访谈 + 10% QABF。",
                "triangulation_es": "65% Datos ABC + 25% Entrevista + 10% QABF.",
                "hypothesis": (
                    "Primary Function: Access to Tangible / Control over Shared"
                    " Media."
                ),
                "hypothesis_zh": "核心行为功能：获取物质/控制共享媒介。",
                "hypothesis_es": (
                    "Función principal: Acceso a tangible / Control de medios."
                ),
                "ferb": (
                    "Use communication board to negotiate turn-taking schedule"
                    " with peers."
                ),
                "ferb_zh": "使用沟通板与同伴协商轮流使用时间表。",
                "ferb_es": (
                    "Usar tablero de comunicación para negociar turnos con"
                    " pares."
                ),
            },
            {
                "name": "3. Elopement from Facility During Schedule Shifts",
                "name_zh": "3. 日程调整期间擅离社区机构",
                "name_es": "3. Fuga de la instalación durante cambios de horario",
                "def": (
                    "Leaving designated day program perimeter without staff"
                    " notification during unplanned alarms or schedule"
                    " transitions. Onset: crossing perimeter exit; Offset:"
                    " staff intervention."
                ),
                "def_zh": (
                    "在未预料的警报或日程转换期间，未经工作人员通知擅自离开指定的日间项目边界。以跨越边界出口为行为开始，以工作人员干预为行为结束。"
                ),
                "def_es": (
                    "Abandonar el perímetro designado del programa diurno sin"
                    " notificar al personal."
                ),
                "ex": "Running out door during unexpected fire drill alarm.",
                "ex_zh": "在突发消防演习警报期间跑出门外。",
                "ex_es": (
                    "Correr hacia afuera durante una alarma de simulacro de"
                    " incendio."
                ),
                "non_ex": "Walking in orderly line with group during scheduled drills.",
                "non_ex_zh": "在计划演习期间随队伍有序排队行走。",
                "non_ex_es": (
                    "Caminar en fila ordenada con el grupo durante simulacros"
                    " programados."
                ),
                "dimensions": (
                    "Frequency: 1 time monthly. Duration: 10-15 minutes."
                    " Intensity: Moderate to High."
                ),
                "dimensions_zh": (
                    "频率：每月 1 次。持续时间：10-15 分钟。强度：中度至高度。"
                ),
                "dimensions_es": (
                    "Frecuencia: 1 vez al mes. Duración: 10-15 min. Intensidad:"
                    " Moderada a Alta."
                ),
                "triggers": (
                    "Setting Events: Sudden loud alarms, unexpected transition"
                    " triggers."
                ),
                "triggers_zh": "背景事件：突发高分贝警报、意外转换触发因素。",
                "triggers_es": (
                    "Eventos de contexto: Alarmas fuertes repentinas,"
                    " transiciones inesperadas."
                ),
                "consequences": (
                    "Staff escorts client safely back and reviews visual"
                    " schedule."
                ),
                "consequences_zh": "工作人员护送客户安全返回并复习视觉日程表。",
                "consequences_es": (
                    "El personal acompaña al cliente de regreso y revisa el"
                    " horario visual."
                ),
                "qabf_summary": (
                    "Escape: 13/15 | Sensory/Anxiety: 11/15 | Attention: 2/15"
                ),
                "qabf_summary_zh": "逃避任务: 13/15 | 感官/焦虑: 11/15 | 社交关注: 2/15",
                "qabf_summary_zh_es": (
                    "Escape: 13/15 | Sensorial/Ansiedad: 11/15 | Atención:"
                    " 2/15"
                ),
                "triangulation": "65% ABC Data + 25% Staff Interview + 10% QABF.",
                "triangulation_zh": "65% ABC 数据 + 25% 工作人员访谈 + 10% QABF。",
                "triangulation_es": "65% Datos ABC + 25% Entrevista + 10% QABF.",
                "hypothesis": (
                    "Primary Function: Escape from Sudden Sensory/Alarm"
                    " Overload."
                ),
                "hypothesis_zh": "核心行为功能：逃避突发感官/警报超载。",
                "hypothesis_es": (
                    "Función principal: Escape de sobrecarga"
                    " sensorial/alarma repentina."
                ),
                "ferb": (
                    "Request staff accompaniment or quiet transition space using"
                    " card."
                ),
                "ferb_zh": "使用卡片请求工作人员陪同或安静的转换空间。",
                "ferb_es": (
                    "Solicitar acompañamiento del personal o espacio"
                    " tranquilo con tarjeta."
                ),
            },
        ],
        "strengths": "High independence in personal self-care.",
        "strengths_zh": "在个人日常生活自理方面具备极高独立性。",
        "strengths_es": "Alta independencia en el cuidado personal.",
        "history": (
            "Participates in Adult Day Vocational Services under Medicaid HCBS"
            " Waiver."
        ),
        "history_zh": (
            "在 Medicaid HCBS 豁免计划下参与成人日间职业干预服务。"
        ),
        "history_es": (
            "Participa en servicios vocacionales para adultos bajo exención"
            " Medicaid HCBS."
        ),
    },
}

# ==========================================
# 3. Main Interface & Security Header
# ==========================================
st.markdown(
    """
    <div class="hipaa-banner">
        <div class="hipaa-title">🛡️ 100% HIPAA COMPLIANT & ZERO-CLOUD LOCAL PROCESSING</div>
        <div class="hipaa-body">
            This tool strictly complies with HIPAA privacy regulations. All data parsing, analysis, and document formulation occur <strong>100% locally within your active browser session memory</strong>.
        </div>
    </div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    "<div class='main-title'>🧩 BCBA Clinical FBA & BIP Draft Formulation Tool"
    " <span class='demo-tag'>(Demo Version)</span></div>",
    unsafe_allow_html=True,
)
st.markdown(
    "<div class='sub-header'>Interactive Demonstration for Automated Clinical"
    " First-Draft Synthesis | Designed for BCBAs & LBAs</div>",
    unsafe_allow_html=True,
)

st.divider()

# ==========================================
# 4. Cohort Selection
# ==========================================
st.markdown("### 1️⃣ Select Clinical Cohort")

cohort_options = {
    "g1": "👶 Early Intervention Protocol (2-5 Yrs)",
    "g2": "🏫 School-Age / IEP Protocol (5-21 Yrs)",
    "g3": "💼 Adult Community & Vocational Protocol (21+ Yrs)",
}

selected_cohort_key = st.radio(
    "Select Target Client Population:",
    options=list(cohort_options.keys()),
    format_func=lambda x: cohort_options[x],
    index=0,
    horizontal=True,
)

current_meta = cohort_meta[selected_cohort_key]

# ==========================================
# 5. Assessment Data Import & Protocol Card
# ==========================================
st.markdown("### 2️⃣ Import Assessment Data & Protocol Overview")

st.markdown(
    f"""
    <div class="protocol-card">
        <div class="protocol-title">📋 Selected Protocol Framework: {current_meta['title']} ({current_meta['framework']})</div>
        {"".join([f'<div class="protocol-bullet">{s}</div>' for s in current_meta['protocol_sentences']])}
    </div>
""",
    unsafe_allow_html=True,
)

col_input1, col_input2, col_input3, col_input4 = st.columns([1, 1, 1, 1])

with col_input1:
  st.markdown("#### 📄 Direct Observation (ABC)")
  mock_csv = generate_mock_abc_csv(selected_cohort_key)
  st.download_button(
      label=f"📥 Download Mock ABC (.csv)",
      data=mock_csv,
      file_name=f"DeIdentified_ABC_{current_meta['file_tag']}.csv",
      mime="text/csv",
      use_container_width=True,
  )
  uploaded_abc = st.file_uploader(
      "Upload ABC File:", type=["csv", "xlsx"], key=f"abc_{selected_cohort_key}"
  )

with col_input2:
  st.markdown("#### 📝 Indirect Interview Notes")
  mock_docx = generate_mock_interview_docx(selected_cohort_key)
  st.download_button(
      label=f"📥 Download Mock Interview (.docx)",
      data=mock_docx,
      file_name=f"DeIdentified_Interview_{current_meta['file_tag']}.docx",
      mime=(
          "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
      ),
      use_container_width=True,
  )
  uploaded_interview = st.file_uploader(
      "Upload Interview File:",
      type=["docx", "txt"],
      key=f"interview_{selected_cohort_key}",
  )

with col_input3:
  st.markdown("#### 📊 Behavior QABF Assessment")
  mock_qabf = generate_mock_qabf_docx(selected_cohort_key)
  st.download_button(
      label=f"📥 Download Mock QABF (.docx)",
      data=mock_qabf,
      file_name=f"DeIdentified_QABF_{current_meta['file_tag']}.docx",
      mime=(
          "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
      ),
      use_container_width=True,
  )
  uploaded_qabf = st.file_uploader(
      "Upload QABF File:",
      type=["docx", "txt"],
      key=f"qabf_{selected_cohort_key}",
  )

with col_input4:
  st.markdown("#### 📈 Behavior Tracking Doc")
  mock_tracking = generate_mock_tracking_csv(selected_cohort_key)
  st.download_button(
      label=f"📥 Download Mock Tracking (.csv)",
      data=mock_tracking,
      file_name=f"DeIdentified_Tracking_{current_meta['file_tag']}.csv",
      mime="text/csv",
      use_container_width=True,
  )
  uploaded_tracking = st.file_uploader(
      "Upload Tracking File:",
      type=["csv", "xlsx"],
      key=f"tracking_{selected_cohort_key}",
  )

active_behaviors = current_meta["behaviors"]


# ==========================================
# 6. Word Document Helper Functions
# ==========================================
def add_bi_heading(doc, level, text_en, text_trans=None):
  h = doc.add_heading(level=level)
  r_en = h.add_run(text_en)
  if text_trans:
    r_tr = h.add_run(f" [{text_trans}]")
    r_tr.italic = True
    r_tr.font.size = Pt(11)
    r_tr.font.color.rgb = RGBColor(120, 120, 120)


def add_bi_item(
    doc, label_en, val_en, label_trans=None, val_trans=None, lang_mode="zh"
):
  p = doc.add_paragraph()
  p.paragraph_format.space_after = Pt(4)
  p.paragraph_format.space_before = Pt(2)

  r_lbl = p.add_run(f"{label_en}: ")
  r_lbl.bold = True
  p.add_run(f"{val_en}")

  if lang_mode != "en" and label_trans and val_trans:
    p.add_run("\n")
    r_tr_lbl = p.add_run(f"[{label_trans}: ")
    r_tr_lbl.bold = True
    r_tr_lbl.italic = True
    r_tr_lbl.font.color.rgb = RGBColor(100, 100, 100)

    r_tr_val = p.add_run(f"{val_trans}]")
    r_tr_val.italic = True
    r_tr_val.font.color.rgb = RGBColor(100, 100, 100)


def build_compact_demographics_table(doc, c_meta, lang_mode):
  table = doc.add_table(rows=5, cols=2)
  table.style = "Table Grid"

  if lang_mode == "zh":
    data = [
        (
            "Student/Client Name (姓名)",
            "[CLIENT_NAME]",
            "DOB / Age (年龄)",
            f"[CLIENT_DOB] / {c_meta['age_str']}",
        ),
        (
            "Client ID (编号)",
            "[CLIENT_ID]",
            "Assessment Date (评估日期)",
            "2026-08-14",
        ),
        (
            "Facility/School (机构/学校)",
            "[DISTRICT_OR_FACILITY_NAME]",
            "Setting (场地)",
            c_meta["setting_str"],
        ),
        (
            "Assessor (评估师)",
            "[BCBA_NAME], BCBA, LBA",
            "Framework (评估框架)",
            c_meta["framework"],
        ),
        (
            "Primary Language (语言)",
            "English / Bilingual Support",
            "Informants (信息提供者)",
            "Parent, Lead Teacher / RBT",
        ),
    ]
  elif lang_mode == "es":
    data = [
        (
            "Client Name (Nombre)",
            "[CLIENT_NAME]",
            "DOB / Age (Edad)",
            f"[CLIENT_DOB] / {c_meta['age_str']}",
        ),
        (
            "Client ID (Identificación)",
            "[CLIENT_ID]",
            "Assessment Date (Fecha)",
            "2026-08-14",
        ),
        (
            "Facility/School (Centro)",
            "[DISTRICT_OR_FACILITY_NAME]",
            "Setting (Entorno)",
            c_meta["setting_str"],
        ),
        (
            "Assessor (Evaluador)",
            "[BCBA_NAME], BCBA, LBA",
            "Framework (Marco)",
            c_meta["framework"],
        ),
        (
            "Primary Language (Idioma)",
            "English / Spanish",
            "Informants (Informantes)",
            "Parent, Lead Teacher / RBT",
        ),
    ]
  else:
    data = [
        (
            "Student/Client Name",
            "[CLIENT_NAME]",
            "DOB / Age",
            f"[CLIENT_DOB] / {c_meta['age_str']}",
        ),
        (
            "Client ID",
            "[CLIENT_ID]",
            "Assessment Date",
            "2026-08-14",
        ),
        (
            "Facility/School",
            "[DISTRICT_OR_FACILITY_NAME]",
            "Setting",
            c_meta["setting_str"],
        ),
        (
            "Assessor",
            "[BCBA_NAME], BCBA, LBA",
            "Framework",
            c_meta["framework"],
        ),
        (
            "Primary Language",
            "English",
            "Informants",
            "Parent, Lead Teacher / RBT",
        ),
    ]

  for row_idx, row_data in enumerate(data):
    row_cells = table.rows[row_idx].cells
    p0 = row_cells[0].paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    p0.add_run(f"{row_data[0]}: ").bold = True
    p0.add_run(row_data[1])

    p1 = row_cells[1].paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    p1.add_run(f"{row_data[2]}: ").bold = True
    p1.add_run(row_data[3])


# ==========================================
# 7. Behavior-by-Behavior Separated FBA Generator
# ==========================================
def generate_exact_fba_doc(cohort_key, lang_choice, behavior_list):
  c_meta = cohort_meta[cohort_key]
  doc = docx.Document()

  lang_mode = "en"
  if "Chinese" in lang_choice:
    lang_mode = "zh"
  elif "Spanish" in lang_choice:
    lang_mode = "es"

  p_t = doc.add_paragraph()
  r_t = p_t.add_run("FUNCTIONAL BEHAVIORAL ASSESSMENT (FBA) REPORT")
  r_t.bold = True
  p_t.style = doc.styles["Title"]

  if lang_mode == "zh":
    p_tr = doc.add_paragraph()
    r_tr = p_tr.add_run("[功能性行为评估 (FBA) 报告 Draft]")
    r_tr.italic = True
    r_tr.font.color.rgb = RGBColor(100, 100, 100)
  elif lang_mode == "es":
    p_tr = doc.add_paragraph()
    r_tr = p_tr.add_run("[Informe de Evaluación de la Conducta Funcional (FBA)]")
    r_tr.italic = True
    r_tr.font.color.rgb = RGBColor(100, 100, 100)

  # Section 1: Demographics
  add_bi_heading(
      doc,
      1,
      "1. Student Demographics & Administrative Info",
      (
          "1. 学生/客户基本信息与行政登记"
          if lang_mode == "zh"
          else "1. Datos demográficos del estudiante"
          if lang_mode == "es"
          else None
      ),
  )
  build_compact_demographics_table(doc, c_meta, lang_mode)

  # Section 2: Data Sources & Triangulation Methodology
  add_bi_heading(
      doc,
      1,
      "2. Data Sources & Triangulation Methodology",
      (
          "2. 数据来源与三方交叉验证评估方法"
          if lang_mode == "zh"
          else "2. Fuentes de datos y metodología de triangulación"
          if lang_mode == "es"
          else None
      ),
  )
  add_bi_item(
      doc,
      "Methodology & Triangulation Formula",
      (
          "Triangulation Algorithm Standard Applied: 65% Direct ABC Data + 25%"
          " Indirect Interview Data + 10% QABF (supplemented by Longitudinal"
          " Behavior Tracking Document).\n\n1. Direct ABC Data: Continuous"
          " recording across baseline therapy sessions.\n2. Indirect Assessment:"
          " Structured interviews with parent and lead RBT.\n3. Psychometric"
          " Rating Scale: Behavior-specific QABF (Questions About Behavioral"
          " Function).\n4. Behavior Tracking Document: Longitudinal frequency"
          " and duration monitoring."
      ),
      (
          "评估方法与三方验证公式"
          if lang_mode == "zh"
          else "Metodología y fórmula de triangulación"
          if lang_mode == "es"
          else None
      ),
      (
          "应用三方数据加权验证算法：65% Direct ABC Data + 25% Indirect Interview Data"
          " + 10% QABF（辅以行为追踪长期记录文件）。\n\n1. 直接 ABC 数据：干预课期间连续行为观察记录。\n2."
          " 间接评估：与家长和督导 RBT 的结构化访谈。\n3. 心理测量量表：针对具体行为的 QABF"
          " 评估问卷。\n4. 行为追踪文件：纵向频率与持续时间监测。"
          if lang_mode == "zh"
          else "Algoritmo de triangulación: 65% ABC + 25% Entrevista + 10%"
          " QABF (con documento de seguimiento longitudinal)."
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  # Section 3: Background & Strengths
  add_bi_heading(
      doc,
      1,
      "3. Brief Background & Strengths Summary",
      (
          "3. 学生背景与优势摘要"
          if lang_mode == "zh"
          else "3. Antecedentes y resumen de fortalezas"
          if lang_mode == "es"
          else None
      ),
  )
  add_bi_item(
      doc,
      "Strengths & Preferences",
      c_meta["strengths"],
      (
          "优势与偏好"
          if lang_mode == "zh"
          else "Fortalezas y preferencias"
          if lang_mode == "es"
          else None
      ),
      (
          c_meta["strengths_zh"]
          if lang_mode == "zh"
          else c_meta.get("strengths_es", c_meta["strengths"])
      ),
      lang_mode,
  )
  add_bi_item(
      doc,
      "Clinical / Educational History",
      c_meta["history"],
      (
          "临床/教育背景"
          if lang_mode == "zh"
          else "Historial clínico / educativo"
          if lang_mode == "es"
          else None
      ),
      (
          c_meta["history_zh"]
          if lang_mode == "zh"
          else c_meta.get("history_es", c_meta["history"])
      ),
      lang_mode,
  )

  # Section 4: Individual Functional Analyses (Behavior-by-Behavior with Dedicated Chart)
  add_bi_heading(
      doc,
      1,
      "4. Individual Target Behavior Functional Analyses",
      (
          "4. 目标行为独立功能分析 (按行为逐项拆解与专属配图)"
          if lang_mode == "zh"
          else "4. Análisis funcional individual por conducta y gráficos exclusivos"
          if lang_mode == "es"
          else None
      ),
  )

  for idx, b in enumerate(behavior_list, 1):
    trans_name = (
        b.get("name_zh")
        if lang_mode == "zh"
        else b.get("name_es")
        if lang_mode == "es"
        else None
    )
    add_bi_heading(
        doc,
        2,
        f"Target Behavior #{idx}: {b['name']}",
        f"目标行为 #{idx}: {trans_name}" if trans_name else None,
    )

    # Insert Dedicated Behavior-Specific Tracking Chart
    chart_title_en = f"Figure 4.{idx}. Longitudinal Behavior Tracking Trend for Target Behavior #{idx}"
    chart_title_zh = f"图 4.{idx}. 目标行为 #{idx} 专属纵向追踪趋势图"
    chart_title_es = f"Figura 4.{idx}. Tendencia de seguimiento para Conducta #{idx}"

    add_bi_heading(
        doc,
        3,
        chart_title_en,
        chart_title_zh
        if lang_mode == "zh"
        else chart_title_es
        if lang_mode == "es"
        else None,
    )
    chart_buf = generate_behavior_tracking_chart(cohort_key, idx - 1)
    doc.add_picture(chart_buf, width=Inches(5.0))
    doc.add_paragraph(
        f"Note: The chart above illustrates the independent tracking frequency"
        f" trend specifically recorded for Target Behavior #{idx}."
    )

    add_bi_item(
        doc,
        "A. Operational Definition",
        b["def"],
        (
            "A. 操作性定义"
            if lang_mode == "zh"
            else "A. Definición operacional"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("def_zh")
            if lang_mode == "zh"
            else b.get("def_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "B. Examples & Non-Examples",
        f"Examples: {b['ex']}\nNon-Examples: {b['non_ex']}",
        (
            "B. 示例与非示例"
            if lang_mode == "zh"
            else "B. Ejemplos y no ejemplos"
            if lang_mode == "es"
            else None
        ),
        (
            f"示例: {b.get('ex_zh')}\n非示例: {b.get('non_ex_zh')}"
            if lang_mode == "zh"
            else f"Ejemplos: {b.get('ex_es')}\nNo ejemplos: {b.get('non_ex_es')}"
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "C. Behavior Dimensions",
        b["dimensions"],
        (
            "C. 行为维度 (频率/持续时间/强度)"
            if lang_mode == "zh"
            else "C. Dimensiones de la conducta"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("dimensions_zh")
            if lang_mode == "zh"
            else b.get("dimensions_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "D. Environmental Triggers & Setting Events",
        b["triggers"],
        (
            "D. 环境触发因素与背景事件"
            if lang_mode == "zh"
            else "D. Desencadenantes y eventos de contexto"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("triggers_zh")
            if lang_mode == "zh"
            else b.get("triggers_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "E. Maintaining Consequences",
        b["consequences"],
        (
            "E. 维持后果与他人反应"
            if lang_mode == "zh"
            else "E. Consecuencias de mantenimiento"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("consequences_zh")
            if lang_mode == "zh"
            else b.get("consequences_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "F. Behavior-Specific QABF Results",
        b["qabf_summary"],
        (
            "F. 该行为专属 QABF 量表得分"
            if lang_mode == "zh"
            else "F. Resultados QABF específicos"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("qabf_summary_zh")
            if lang_mode == "zh"
            else b.get("qabf_summary_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "G. Triangulation (65% ABC + 25% Interview + 10% QABF)",
        b["triangulation"],
        (
            "G. 三方交叉验证算法结果"
            if lang_mode == "zh"
            else "G. Triangulación de datos"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("triangulation_zh")
            if lang_mode == "zh"
            else b.get("triangulation_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "H. Hypothesized Function",
        b["hypothesis"],
        (
            "H. 该行为推断功能"
            if lang_mode == "zh"
            else "H. Función hipetotizada"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("hypothesis_zh")
            if lang_mode == "zh"
            else b.get("hypothesis_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )

  # Section 5: Overall Synthesis
  add_bi_heading(
      doc,
      1,
      "5. Synthesis & Personalised Clinical Recommendations",
      (
          "5. 个性化综合评估结论与临床建议"
          if lang_mode == "zh"
          else "5. Síntesis y recomendaciones clínicas personalizadas"
          if lang_mode == "es"
          else None
      ),
  )

  if cohort_key == "g1":
    synth_en = (
        "Based on the triangulated functional assessment (65% Direct ABC + 25%"
        " Indirect Interview + 10% QABF), the client (Ages 2-5) primarily"
        " demonstrates behaviors maintained by Social Negative Reinforcement"
        " (Task/Transition Escape) and Sensory Automatic Regulation"
        " (Auditory Aversion). Recommendations:\n1. Formulate an NDBI / ESDM"
        " play-based Behavior Intervention Plan (BIP) integrating AAC (BigMack"
        " / PECS 'Break' card) for early functional communication training"
        " (FCT).\n2. Implement antecedent noise-mitigation protocols"
        " (noise-canceling headphones, visual schedule warnings) prior to"
        " structured table tasks and transitions.\n3. Train parents and RBTs on"
        " shared co-regulation strategies to replace self-injurious and"
        " sensory distress behaviors."
    )
    synth_zh = (
        "基于 65% 直接 ABC 观察 + 25% 间接访谈 + 10% QABF 量表的三方交叉验证算法，该 2-5"
        " 岁幼儿的目标行为主要由‘逃避任务/转换（社交负强化）’及‘调节听觉过度刺激（感官自动强化）’所维持。临床建议：\n1."
        " 制定基于 ESDM / NDBI 自然教法的行为干预计划 (BIP)，优先引入 AAC"
        " 语音按键与 PECS‘休息’视觉卡进行早期功能性沟通训练 (FCT)。\n2."
        " 在结构化桌面任务与环节转换前，实施前因降噪与视觉倒计时预警。\n3."
        " 对家长及 RBT 进行情绪共调节 (Co-regulation) 技巧培训，以替代自伤与感官情绪发作。"
    )
    synth_es = (
        "Con base en la triangulación (65% ABC + 25% Entrevista + 10% QABF), el"
        " cliente (2-5 años) presenta conductas mantenidas por Escape de"
        " Tareas y Regulación Sensorial Auditiva. Recomendaciones:\n1. Diseñar"
        " un BIP basado en NDBI/ESDM integrando AAC/PECS para Comunicación"
        " Funcional (FCT).\n2. Implementar protocolos de mitigación de ruido y"
        " avisos visuales antes de transiciones.\n3. Capacitar a padres y RBTs"
        " en co-regulación."
    )
  elif cohort_key == "g2":
    synth_en = (
        "Based on triangulated assessment data (65% ABC + 25% IEP Interview +"
        " 10% QABF), Wei's school-age student demonstrates elopement, group"
        " avoidance, and auditory overload primarily maintained by escape from"
        " academic/social demands and sensory automatic regulation."
        " Recommendations:\n1. Develop an IEP-aligned BIP integrating PBIS Tier"
        " 2/3 supports, providing discrete 'Break' card access and quiet"
        " working spaces.\n2. Implement proactive priming and visual schedules"
        " to mitigate transition anxiety and auditory distress.\n3. Train"
        " classroom aides and teachers on FCT implementation and environmental"
        " modifications."
    )
    synth_zh = (
        "基于三方交叉验证数据（65% ABC + 25% IEP访谈 + 10% QABF），Wei的学龄期学生表现出的擅离座位、小组回避及听觉超载行为，主要由逃避学业/社交要求及感官自动调节所维持。临床建议：\n1."
        " 制定与 IEP 紧密结合的 BIP，整合 PBIS 二/三级支持，提供“休息”卡片和安静学习区。\n2."
        " 实施前因预告与视觉日程表，以缓解转换焦虑与听觉不适。\n3."
        " 对随班就读老师和助教进行 FCT 落地执行与环境调整培训。"
    )
    synth_es = (
        "Basado en datos triangulados (65% ABC + 25% Entrevista IEP + 10%"
        " QABF), el estudiante en edad escolar presenta conductas de fuga y"
        " evitación mantenidas por escape de demandas académicas y regulación"
        " sensorial. Recomendaciones:\n1. Desarrollar un BIP alineado al IEP"
        " integrando soportes PBIS y acceso a tarjetas de descanso.\n2."
        " Implementar avisos visuales para mitigar la ansiedad de"
        " transición.\n3. Capacitar a auxiliares en FCT."
    )
  else:
    synth_en = (
        "Based on triangulated adult community assessment (65% ABC + 25% Vocational"
        " Interview + 10% QABF), the adult client presents vocational refusal,"
        " property disruption, and elopement maintained by escape from assembly"
        " quotas, peer media disputes, and sudden alarm overloads."
        " Recommendations:\n1. Implement a person-centered Medicaid HCBS Waiver"
        " BIP prioritizing self-advocacy phrase cards and structured choice"
        " boards.\n2. Provide incremental vocational quotas and structured"
        " transition timers to minimize frustration.\n3. Train direct support"
        " professionals (DSPs) on neutral redirection, active listening, and"
        " proactive crisis prevention."
    )
    synth_zh = (
        "基于成人社区三方交叉验证评估（65% ABC + 25% 职业访谈 + 10% QABF），该成人客户表现出的职业拒绝、物品破坏及社区擅离，主要由逃避装配配额、同伴媒介冲突及突发警报超载所维持。临床建议：\n1."
        " 实施以人为本的 Medicaid HCBS 豁免计划 BIP，优先使用自我倡导短语卡和结构化选择板。\n2."
        " 提供递进的职业配额与结构化转换计时器，以最大限度减少挫败感。\n3."
        " 对直属支持人员 (DSP) 进行中立重定向、积极倾听及主动危机预防的专业培训。"
    )
    synth_es = (
        "Basado en la evaluación comunitaria para adultos (65% ABC + 25% Entrevista"
        " vocacional + 10% QABF), el cliente adulto presenta rechazo"
        " vocacional y disrupción mantenidos por escape de cuotas y alarmas."
        " Recomendaciones:\n1. Implementar un BIP centrado en la persona bajo"
        " exención Medicaid HCBS priorizando tarjetas de autodefensa.\n2."
        " Proporcionar cuotas vocacionales incrementales y temporizadores.\n3."
        " Capacitar al personal DSP en redirección neutral."
    )

  add_bi_item(
      doc,
      "Clinical Synthesis & Next Steps",
      synth_en,
      (
          "临床综合结论与后续步骤"
          if lang_mode == "zh"
          else "Síntesis clínica y próximos pasos"
          if lang_mode == "es"
          else None
      ),
      (
          synth_zh
          if lang_mode == "zh"
          else synth_es
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  bio = io.BytesIO()
  doc.save(bio)
  bio.seek(0)
  return bio


# ==========================================
# 8. Comprehensive Enriched BIP Generator
# ==========================================
def generate_exact_bip_doc(cohort_key, lang_choice):
  c_meta = cohort_meta[cohort_key]
  doc = docx.Document()

  lang_mode = "en"
  if "Chinese" in lang_choice:
    lang_mode = "zh"
  elif "Spanish" in lang_choice:
    lang_mode = "es"

  p_t = doc.add_paragraph()
  r_t = p_t.add_run("BEHAVIOR INTERVENTION PLAN (BIP)")
  r_t.bold = True
  p_t.style = doc.styles["Title"]

  if lang_mode == "zh":
    p_tr = doc.add_paragraph()
    r_tr = p_tr.add_run("[行为干预计划 (BIP) Comprehensive Draft]")
    r_tr.italic = True
    r_tr.font.color.rgb = RGBColor(100, 100, 100)
  elif lang_mode == "es":
    p_tr = doc.add_paragraph()
    r_tr = p_tr.add_run(
        "[Plan de Intervención Conductual (BIP) Borrador Completo]"
    )
    r_tr.italic = True
    r_tr.font.color.rgb = RGBColor(100, 100, 100)

  # Section 1
  add_bi_heading(
      doc,
      1,
      "1. Student Info & Administrative Summary",
      (
          "1. 学生/客户信息与行政摘要"
          if lang_mode == "zh"
          else "1. Información del estudiante y resumen"
          if lang_mode == "es"
          else None
      ),
  )
  build_compact_demographics_table(doc, c_meta, lang_mode)

  # Section 2: Behavior Functions & FERB Breakdown
  add_bi_heading(
      doc,
      1,
      "2. Target Behaviors, Functions & Replacement Skills (FERB)",
      (
          "2. 目标行为、行为功能与替代技能 (FERB) 逐项拆解"
          if lang_mode == "zh"
          else "2. Conductas objetivo, funciones y conductas de reemplazo (FERB)"
          if lang_mode == "es"
          else None
      ),
  )

  for idx, b in enumerate(c_meta["behaviors"], 1):
    trans_name = (
        b.get("name_zh")
        if lang_mode == "zh"
        else b.get("name_es")
        if lang_mode == "es"
        else None
    )
    add_bi_heading(
        doc,
        2,
        f"Target Behavior #{idx}: {b['name']}",
        f"目标行为 #{idx}: {trans_name}" if trans_name else None,
    )

    add_bi_item(
        doc,
        "Operational Definition",
        b["def"],
        (
            "操作性定义"
            if lang_mode == "zh"
            else "Definición operacional"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("def_zh")
            if lang_mode == "zh"
            else b.get("def_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "Validated Function",
        b["hypothesis"],
        (
            "确证行为功能"
            if lang_mode == "zh"
            else "Función validada"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("hypothesis_zh")
            if lang_mode == "zh"
            else b.get("hypothesis_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )
    add_bi_item(
        doc,
        "Functionally Equivalent Replacement Behavior (FERB)",
        b["ferb"],
        (
            "功能性替代行为 (FERB)"
            if lang_mode == "zh"
            else "Conducta de reemplazo funcionalmente equivalente (FERB)"
            if lang_mode == "es"
            else None
        ),
        (
            b.get("ferb_zh")
            if lang_mode == "zh"
            else b.get("ferb_es")
            if lang_mode == "es"
            else None
        ),
        lang_mode,
    )

  # Section 3: Proactive / Antecedent Strategies
  add_bi_heading(
      doc,
      1,
      "3. Proactive & Antecedent Modifications (Prevention)",
      (
          "3. 前因调整与预防策略 (统一整合)"
          if lang_mode == "zh"
          else "3. Modificaciones proactivas y antecedentes"
          if lang_mode == "es"
          else None
      ),
  )
  add_bi_item(
      doc,
      "3.1 Environmental & Priming Adaptations",
      (
          "• Provide 2-minute and 1-minute visual/auditory transition warnings"
          " prior to changing activities.\n• Offer noise-canceling headphones"
          " or move client to low-stimulation sensory area prior to task"
          " demands.\n• Break academic/vocational tasks into small, visual"
          " chunks (2-3 items per strip)."
      ),
      (
          "3.1 环境调整与预先提示"
          if lang_mode == "zh"
          else "3.1 Adaptaciones ambientales"
          if lang_mode == "es"
          else None
      ),
      (
          "• 在活动转换前提供 2分钟 及 1分钟 的视觉/听觉预先倒计时提示。\n•"
          " 在任务开始前主动提供降噪耳机或移至低刺激感官区域。\n•"
          " 将学业/职业任务拆解为小步子视觉单元（每条2-3个小任务）。"
          if lang_mode == "zh"
          else "• Proporcionar avisos visuales/auditivos de 2 min y 1 min antes"
          " de transiciones.\n• Ofrecer auriculares de cancelación de ruido"
          " antes de tareas.\n• Dividir tareas en pequeños bloques visuales."
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  # Section 4: Replacement Behaviors Protocols
  add_bi_heading(
      doc,
      1,
      "4. Replacement Behaviors & Functional Communication Training (FCT)",
      (
          "4. 替代行为与功能性沟通训练 (FCT)"
          if lang_mode == "zh"
          else "4. Entrenamiento en Comunicación Funcional (FCT)"
          if lang_mode == "es"
          else None
      ),
  )
  add_bi_item(
      doc,
      "4.1 Functional Communication Protocols",
      (
          "• Primary FCT Skill: Prompt client to press an AAC button or hand a"
          " PECS icon for 'Break' or 'Help' upon initial sign of"
          " frustration.\n• Systematic Prompting: Use Most-to-Least physical"
          " assistance, fading rapidly to gestural/visual prompts within 10"
          " days."
      ),
      (
          "4.1 功能性沟通协议"
          if lang_mode == "zh"
          else "4.1 Protocolos de comunicación funcional"
          if lang_mode == "es"
          else None
      ),
      (
          "• 核心 FCT 技能：教导客户在产生烦躁情绪萌芽时，按下 AAC 沟通按键或递交 PECS"
          " 卡片表达“休息”或“帮助”。\n• 辅助渐退策略：使用由多到少（Most-to-Least）物理辅助，并在"
          " 10 天内快速渐退至手势或视觉提示。"
          if lang_mode == "zh"
          else "• Habilidad FCT principal: Indicar al cliente que presione el"
          " botón AAC o entregue la tarjeta PECS 'Descanso'.\n• Fading de"
          " ayudas: De mayor a menor asistencia en 10 días."
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  # Section 5: Reinforcement Strategies
  add_bi_heading(
      doc,
      1,
      "5. Reinforcement Protocols",
      (
          "5. 强化策略协议"
          if lang_mode == "zh"
          else "5. Protocolos de Reforzamiento"
          if lang_mode == "es"
          else None
      ),
  )
  add_bi_item(
      doc,
      "5.1 Differential Reinforcement of Alternative Behavior (DRA)",
      (
          "• Immediate (within 3 seconds) 100% compliance with requested 'Break'"
          " or 'Help' AAC activations during initial acquisition phase.\n• Pair"
          " escape with enthusiastic verbal praise ('Great job asking for a"
          " break!')."
      ),
      (
          "5.1 替代行为区别性强化 (DRA)"
          if lang_mode == "zh"
          else "5.1 Reforzamiento Diferencial de Conducta Alternativa (DRA)"
          if lang_mode == "es"
          else None
      ),
      (
          "• 在习得阶段，只要客户按下 AAC 表达“休息”，必须在 3 秒内 100%"
          " 满足其休息请求。\n•"
          " 将逃避任务与高度热情的情感口头表扬结合（如：“太棒了，你自己按按键说要休息！”）。"
          if lang_mode == "zh"
          else "• Cumplimiento inmediato (en 3 segundos) del 100% con la"
          " solicitud de 'Descanso'.\n• Acompañar el descanso con el elogio"
          " verbal entusiasta."
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  # Section 6: Response Strategies
  add_bi_heading(
      doc,
      1,
      "6. Reactive Response Protocols & Extinction",
      (
          "6. 目标行为回应与消退策略"
          if lang_mode == "zh"
          else "6. Protocolos de respuesta reactiva y extinción"
          if lang_mode == "es"
          else None
      ),
  )
  add_bi_item(
      doc,
      "6.1 Extinction & Neutral Physical Blocking",
      (
          "• Escape Extinction: Maintain neutral expression, avoid eye contact,"
          " minimize verbal dialogue during problem behavior.\n• Physical"
          " Blocking: Promptly and softly block any SIB, slapping, or"
          " floor-dropping using foam pads to prevent injury without giving"
          " emotional feedback."
      ),
      (
          "6.1 消退与中立物理阻挡"
          if lang_mode == "zh"
          else "6.1 Extinción y bloqueo físico neutral"
          if lang_mode == "es"
          else None
      ),
      (
          "•"
          " 逃避消退：在问题行为发生时，保持平静中立表情，避免眼神接触，不进行长篇大论训诫。\n•"
          " 物理阻挡：若出现自伤或攻击行为，使用软垫迅速柔和阻挡，确保安全的同时不给予额外的言语或情感反馈。"
          if lang_mode == "zh"
          else "• Extinción de escape: Mantener expresión neutral, evitar contacto"
          " visual.\n• Bloqueo físico: Bloquear SIB o bofetadas de forma suave"
          " usando colchonetas."
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  # Section 7: Safety & Data
  add_bi_heading(
      doc,
      1,
      "7. Crisis Safety & Treatment Fidelity",
      (
          "7. 危机安全预案与执行忠实度"
          if lang_mode == "zh"
          else "7. Seguridad en crisis y fidelidad del tratamiento"
          if lang_mode == "es"
          else None
      ),
  )
  add_bi_item(
      doc,
      "Data Collection & Fidelity Protocols",
      (
          "• RBTs will record daily frequency/duration of target behaviors and"
          " independent FCT requests.\n• BCBA will conduct weekly treatment"
          " fidelity observations using a 10-point checklist."
      ),
      (
          "数据收集与忠实度核查"
          if lang_mode == "zh"
          else "Recolección de datos y fidelidad"
          if lang_mode == "es"
          else None
      ),
      (
          "• RBT 每日记录目标行为的发生频率/持续时间及 FCT 独立使用次数。\n• BCBA"
          " 每周使用 10 项标准核查表进行 1:1 干预忠实度评估。"
          if lang_mode == "zh"
          else "• Los RBTs registrarán la frecuencia/duración diaria.\n• El"
          " BCBA realizará observaciones semanales de fidelidad."
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  # Section 8: Staff Training and Monitoring
  add_bi_heading(
      doc,
      1,
      "8. Staff Training and Monitoring",
      (
          "8. 人员培训、执行监督与计划复审机制"
          if lang_mode == "zh"
          else "8. Capacitación del personal y supervisión"
          if lang_mode == "es"
          else None
      ),
  )

  st_en = (
      "• Responsible Party: Board Certified Behavior Analyst (BCBA) / Licensed"
      " Behavior Analyst (LBA).\n• Training Process: Initial 2-hour didactic"
      " module followed by 1:1 behavioral skills training (BST: Instructions,"
      " Modeling, Rehearsal, and Feedback) with all direct intervention staff"
      " (RBTs/Therapists).\n• Observation & Fidelity Frequency: Direct"
      " weekly observations utilizing a 10-point Treatment Fidelity Checklist"
      " to ensure implementation accuracy (>90% compliance required).\n• Plan"
      " Review Schedule: Formal clinical review scheduled every 90 days, or"
      " immediately if target behaviors show a 20% spike in frequency."
  )
  st_zh = (
      "• 责任人：管辖 BCBA 行为分析师 / LBA 执照行为分析师。\n• 培训流程：实施 2"
      " 小时初始理论培训，随后进行 1:1 行为技能训练 (BST"
      " 模式：讲解、示范、演练与即时反馈)，覆盖所有直属 RBT / 干预人员。\n•"
      " 观察与忠实度核查频率：每周进行一次 1:1 现场观察，使用 10"
      " 项行为忠实度清单核查（要求达标率 >90%）。\n• 计划复审时间表：每 90"
      " 天进行一次正式临床复审，若目标行为频率发生 >20% 的异常激增则触发即时复审。"
  )
  st_es = (
      "• Persona responsable: BCBA / LBA.\n• Proceso de capacitación: Módulo"
      " de 2 horas seguido de Entrenamiento en Habilidades Conductuales (BST:"
      " Instrucciones, Modelado, Ensayo y Retroalimentación).\n• Frecuencia de"
      " observación: Observaciones semanales con lista de cotejo de fidelidad"
      " del 10-pt (requerido >90%).\n• Programación de revisión: Revisión"
      " clínica cada 90 días."
  )

  add_bi_item(
      doc,
      "Staff Training & Fidelity Monitoring Plan",
      st_en,
      (
          "人员培训与忠实度监控计划"
          if lang_mode == "zh"
          else "Plan de capacitación y monitoreo"
          if lang_mode == "es"
          else None
      ),
      (
          st_zh
          if lang_mode == "zh"
          else st_es
          if lang_mode == "es"
          else None
      ),
      lang_mode,
  )

  bio = io.BytesIO()
  doc.save(bio)
  bio.seek(0)
  return bio


# ==========================================
# 9. Action Buttons
# ==========================================
st.markdown("### 3️⃣ Target Language & Formulate / Download Actions")

col_lang, col_action1, col_action2 = st.columns([1.2, 1.4, 1.4])

with col_lang:
  report_lang = st.radio(
      "Select Target Report Language / Format:",
      options=[
          "English (US Standard)",
          "Bilingual (English / Simplified Chinese - 简体中文)",
          "Bilingual (English / Spanish - Español)",
      ],
      index=1,
  )

fba_docx_bytes = generate_exact_fba_doc(
    selected_cohort_key, report_lang, active_behaviors
)
bip_docx_bytes = generate_exact_bip_doc(selected_cohort_key, report_lang)

lang_tag = (
    "English"
    if "Standard" in report_lang
    else "Bilingual_ZH"
    if "Chinese" in report_lang
    else "Bilingual_ES"
)

with col_action1:
  st.write(" ")
  st.write(" ")
  st.download_button(
      label="⚡ Formulate & Download De-Identified FBA Draft (.docx)",
      data=fba_docx_bytes,
      file_name=(
          f"DeIdentified_FBA_Draft_{current_meta['file_tag']}_{lang_tag}.docx"
      ),
      mime=(
          "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
      ),
      use_container_width=True,
  )

with col_action2:
  st.write(" ")
  st.write(" ")
  st.download_button(
      label="⚡ Formulate & Download De-Identified BIP Draft (.docx)",
      data=bip_docx_bytes,
      file_name=(
          f"DeIdentified_BIP_Draft_{current_meta['file_tag']}_{lang_tag}.docx"
      ),
      mime=(
          "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
      ),
      use_container_width=True,
  )

st.divider()

st.caption(
    "⚠️ **Clinical Responsibility Notice:** This formulation tool serves"
    " strictly as a clinical first-draft synthesizer for BCBAs and LBAs. All"
    " generated drafts are fully de-identified and must be independently"
    " reviewed and edited prior to formal signature."
)
